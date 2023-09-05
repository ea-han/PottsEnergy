import os
import json
import pandas as pd
import csv
from docx import Document
from docx.shared import Pt

##This class imports mutations-energy pairs and classifies them into rescue, antagonistic and compensatory.
##

class ClassifySequences:
    seqStack = []
    consensus = []
    jFileArray = []
    reduxName = ""
    mutationList = []

    
    INTERACTION1 = "Rescue"
    INTERACTION2 = "Compensatory"
    INTERACTION3 = "Antagonistic"

    ##search list of sequences for mutations 1 and 2
    ##add found mutations to a list of sequences to calculate energy of
    ##reduce sequences and mutations with redux file
    ##calculate average of all sequences
    ##mutations: [index, source, [mut1, mut2,etc]],[index2, source2, [mut1, mut2,etc]],

    def __init__(self,seqStack, consensus, mutationList, jFileArray, reduxName):
        self.foundSeqs = []
        self.seqStack = seqStack
        self.reduxName = reduxName
        self.jFileArray = jFileArray
        self.consensus = self.reduceSequence(consensus)
        self.mutationList = mutationList

        mutList = []
        for mut in mutationList:
            mutList.append(self.mutPairToString(mut))

        with open("src/mutationNames.json","w") as ifile:
            json.dump(mutList,ifile)


    ## get energy for consensus methods {
    def getCoupling(self, iVal,jVal):
        aVali = ord(iVal) - 65
        aValj = ord(jVal) - 65
        if (aVali < 0 or aVali > 3):
            print("invalid coupling")
            return
        if (aVali == 0 and aValj == 0):
            return 0
        if (aVali == 0 and aValj == 1):
            return 1
        if (aVali == 0 and aValj == 2):
            return 2
        if (aVali == 0 and aValj == 3):
            return 3
        if (aVali == 1 and aValj == 0):
            return 4
        if (aVali == 1 and aValj == 1):
            return 5
        if (aVali == 1 and aValj == 2):
            return 6
        if (aVali == 1 and aValj == 3):
            return 7
        if (aVali == 2 and aValj == 0):
            return 8
        if (aVali == 2 and aValj == 1):
            return 9
        if (aVali == 2 and aValj == 2):
            return 10
        if (aVali == 2 and aValj == 3):
            return 11
        if (aVali == 3 and aValj == 0):
            return 12
        if (aVali == 3 and aValj == 1):
            return 13
        if (aVali == 3 and aValj == 2):
            return 14
        if (aVali == 3 and aValj == 3):
            return 15
        return -1

    ##gets energy of a sequence
    def getEnergy(self, seq):
        ctr = 0
        energy = 0
        length = len(seq)

        for i in range(0,length):
            iAcid = seq[i]
            for j in range (i+1,length):
                jAcid = seq[j]
                couplingIdx = self.getCoupling(iAcid,jAcid)
                if couplingIdx < 0:
                    print("coupling failed")
                    return -1
                energy += self.jFileArray[ctr][couplingIdx]
                ctr+=1
        return energy

    def mutateSequence1Consensus(self, seq,mutation):
        sequence = list(seq)
        sequence[mutation[0]] = mutation[2]
        return sequence
    
    def mutateSequence2Consensus(self, seq,mutation1,mutation2):
        sequence = list(seq)
        sequence[mutation1[0]] = mutation1[2]
        sequence[mutation2[0]] = mutation2[2]
        return sequence
    
    ##takes reduced sequence andmutation and calcultes energies. Need to screen for valid mutation/sequences.
    #runs? need to check avlues
    def calcEnergySequence(self,seq,name,mutationArray):
        
        defaultEnergy = self.getEnergy(seq) ##works
        m1Energy = self.getEnergy(self.mutateSequence1Consensus(seq, mutationArray[0]))
        m2Energy = self.getEnergy(self.mutateSequence1Consensus(seq, mutationArray[1]))
        m1m2Energy = self.getEnergy(self.mutateSequence2Consensus(seq,mutationArray[0],mutationArray[1])) ##works

        m1Delta = defaultEnergy - m1Energy 
        m2Delta = defaultEnergy - m2Energy
        m1m2Delta = defaultEnergy - m1m2Energy 
        deltaDeltaE = m1m2Delta - (m1Delta + m2Delta)

        return [mutationArray,name,deltaDeltaE,m1m2Delta,m1Delta,m2Delta]

    #}



    def arrToString(self,mutation):
        altMutations = ','.join(mutation[2])
        newString = ''
        newString = newString + str(mutation[1])
        newString = newString + str(mutation[0]+1)
        newString = newString + altMutations
        return newString

    def mutPairToString(self,mutationPair):
        string1 = self.arrToString(mutationPair[0])
        string2 = self.arrToString(mutationPair[1])
        return string1 + "-" + string2

    def reduceSequence(self, seq):
        ctr = 0
        with open(self.reduxName) as kFile:
                returnSeq = []
                for line in kFile:
                    ##simlify
                    line = line.strip()
                    insertArr = line.split()
                    ctr = int(insertArr[0])
                    aAlpha = list(insertArr[1])
                    bAlpha = list(insertArr[2])
                    cAlpha = list(insertArr[3])
                    dAlpha = list(insertArr[4])
                    
                    val = seq[ctr]
                    if val in aAlpha:
                        returnSeq.append('A')
                    elif val in bAlpha:
                        returnSeq.append('B')
                    elif val in cAlpha:
                        returnSeq.append('C')
                    elif val in dAlpha:
                        returnSeq.append('D')
                    else:
                        print("`invalid char`:",val ,flush=True)
                        return -1
                    ctr+=1
        return returnSeq


    ##takes in SINGLE MUTATION valuestack,compares each sequence DDE with consensus and classifies.
    ##Sorts and returns all values in valuestack
    ##returns []
    ##REWRITE
    def getInteractionsAndDifference(self,mutationData):
        mutPair = mutationData[0]
        valueStack = mutationData[1]

        rescue = []
        compensate = []
        antag = []

        ##pdb.set_trace()

        #get consensus interaction
        consensusData = self.calcEnergySequence(self.consensus,"consensus", mutPair)
        consensusInteraction = self.getInteraction(consensusData[4],consensusData[5],consensusData[3])
        print(consensusData, consensusInteraction)

        ##mutationpair, seq index, delta delta, d12, d1, d2

        index = 0
        for seqArr in valueStack:
            seq = self.seqStack[seqArr[1]]
            actualDeltam1m2 = seqArr[3]
            Dm1 = seqArr[4]
            Dm2 = seqArr[5]
            DDe = seqArr[2]
            hd = self.getHDNum(seq)
            seqMutList =  len(hd)

            comparisonData = self.getComparisonDelta(Dm1,Dm2,actualDeltam1m2)

            distFromCompValue = comparisonData[0]
            isRescue = comparisonData[1]

            ##check interaction
            ##seqData = [mutPair,index,interaction,distFromMaxDelta1,actualDeltam1m2]
            seqData = [Dm1,Dm2,actualDeltam1m2,DDe,hd,seqMutList,distFromCompValue]


            ##new comparison func
            index+=1
            if isRescue:
                rescue.append(seqData)
            elif distFromCompValue >= 0:
                compensate.append(seqData)
            else:
                antag.append(seqData)
        
        return [rescue,compensate,antag,consensusData,consensusInteraction]


    def getHDNum(self,seq):
        seqMut = ''
        consMut = ''
        mutList = []
        for i in range(0,len(self.consensus)):
            seqMut = seq[i] 
            consMut = self.consensus[i]
            if seqMut != consMut:
                addString = consMut+str(i)+seqMut
                mutList.append(addString)

        return mutList


    ##gets comparison data. if rescue, gets max distance from top mutation. Else, gets distance from bottom mutation.
    def getComparisonDelta(self,Dm1,Dm2,Dm1m2):
        compVal1 = max(Dm1, Dm2)
        compVal2 = min(Dm1, Dm2)
        isRescue = False
        if Dm1m2 > compVal1:
            delta = Dm1m2 - compVal1
            isRescue = True
        else:
            delta = Dm1m2 - compVal2
        return [delta,isRescue]
    
    def importValueStacks(self,filename):
        with open(filename) as oFile:
            return json.load(oFile)


    #returns interaction
    def getInteraction(self,Dm1,Dm2, actual):
        projectedDelta = Dm1 + Dm2
        if actual < projectedDelta:
            return self.INTERACTION3
        elif actual > max(Dm1,Dm2):
            return self.INTERACTION1
        else:
            return self.INTERACTION2
   
    def makeFlippedList(self,dataStack, consData):
        flippedList = []
        ##checks if consensus m1 is greater than or equal to consensus m2
        consensusm1Gm2 = consData[0] >= consData[1]
 
        for seqData in dataStack:
            ##gets isPositve of m1-m2 for comparison
            sequencem1Gm2 = seqData[0] >= seqData[1]

            ##if the m1 IsGreaterThan m2 boolean is different then the roles of the mutations m1 m2 must be swapped
            ## append to flippedlist
            if consensusm1Gm2 != sequencem1Gm2:
                flippedList.append(seqData)

        return flippedList


    ##compares value stack interactions with consensus. sorts into one big list by interaction, delta difference.
    ##combine all relevant data into one entry. 
    ## label,index,interaction,consensusInteraction,deltaDifference
    ##returns 2 arrays:[same as consensus, diff to consensus]
    def compareToConsensus(self):
        ##gets values: [[mutation,computed stack values],[mutation,computed stack values],...]
        sequenceCalculations = self.importValueStacks("src/allValueStacks.json")



        mutationListCtr = 0
        for mutationData in sequenceCalculations:

            print(mutationData)
            mutName = self.mutPairToString(mutationData[0])
            print("mutname = ", mutName)

            seqData = self.getInteractionsAndDifference(mutationData)
            rescue = seqData[0]
            compensate = seqData[1]
            antag = seqData[2]
            

            #sorts
            rescue = sorted(rescue,key=lambda x: x[6],reverse=True)
            compensate = sorted(compensate,key=lambda x: x[6],reverse=True)
            antag = sorted(antag,key=lambda x: x[6],reverse=True)

            print("computed for mutation ",mutName)
            print("sorted by distance from Dm1m2 to highest-fitness single mutation")
            ##             seqData = [Dm1,Dm2,actualDeltam1m2,DDe,hd,seqMutList,distFromCompValue]


            # [mutationArray,name,deltaDeltaE,m1m2Delta,m1Delta,m2Delta]
            consensusRow = seqData[3]
            consensusInteraction = seqData[4]
            lastRow = [consensusRow[4],consensusRow[5],consensusRow[3],consensusRow[2],"None",0,0]


            if (consensusInteraction == self.INTERACTION1):
                flippedList = self.makeFlippedList(rescue,lastRow)
            elif(consensusInteraction == self.INTERACTION2):
                flippedList = self.makeFlippedList(compensate,lastRow)
            else:
                flippedList = self.makeFlippedList(antag,lastRow)



            rescue.insert(0,lastRow)
            compensate.insert(0,lastRow)
            antag.insert(0,lastRow)
            flippedList.insert(0,lastRow)




            if (not os.path.exists("out/"+mutName)):
                os.mkdir("out/"+mutName)

            #writes files
            self.writeToJson(rescue,"out/"+mutName+"/rescue")
            self.writeToJson(compensate,"out/"+mutName+"/compensate")            
            self.writeToJson(antag,"out/"+mutName+"/antag")
            self.writeToJson(flippedList,"out/"+mutName+"/flipped")

            self.writeToCSV(rescue,"out/"+mutName+"/rescue")
            self.writeToCSV(compensate,"out/"+mutName+"/compensate")
            self.writeToCSV(antag,"out/"+mutName+"/antag")
            self.writeToCSV(flippedList,"out/"+mutName+"/flipped")

            self.writeToDocx(rescue,compensate,antag,flippedList,"out/"+mutName)

            ##updates index of original mutations
            mutationListCtr +=1

    ##saves top 5 of antag,comp,rescue,flipped of mutation to docx file.
    def writeToDocx(self,rescue,compensate,antag,flipped,name):
        ##slicing, inserting header
        fields = ["Dm1","Dm2","Dm1m2","DDe","list of mutations","HD","distance from max/min D single mutation"]
        

        rescue = rescue[0:5]
        compensate = compensate[0:5]
        antag = antag[0:5]
        flipped = flipped[0:5]


        ##make dataframes
        rescueFrame = pd.DataFrame(data=rescue,columns=fields)
        compensateFrame = pd.DataFrame(data=compensate,columns=fields)
        antagFrame = pd.DataFrame(data=antag,columns=fields)
        flipFrame = pd.DataFrame(data=flipped,columns=fields)
        
        rescueFrame['Dm1'] = rescueFrame['Dm1'].astype(float).round(4)
        rescueFrame['Dm2'] = rescueFrame['Dm2'].astype(float).round(4)
        rescueFrame['Dm1m2'] = rescueFrame['Dm1m2'].astype(float).round(4)
        rescueFrame['DDe'] = rescueFrame['DDe'].astype(float).round(4)
        rescueFrame["distance from max/min D single mutation"] = rescueFrame["distance from max/min D single mutation"].astype(float).round(4)

        antagFrame['Dm1'] = antagFrame['Dm1'].astype(float).round(4)
        antagFrame['Dm2'] = antagFrame['Dm2'].astype(float).round(4)
        antagFrame['Dm1m2'] = antagFrame['Dm1m2'].astype(float).round(4)
        antagFrame['DDe'] = antagFrame['DDe'].astype(float).round(4)
        antagFrame["distance from max/min D single mutation"] = antagFrame["distance from max/min D single mutation"].astype(float).round(4)


        compensateFrame['Dm1'] = compensateFrame['Dm1'].astype(float).round(4)
        compensateFrame['Dm2'] = compensateFrame['Dm2'].astype(float).round(4)
        compensateFrame['Dm1m2'] = compensateFrame['Dm1m2'].astype(float).round(4)
        compensateFrame['DDe'] = compensateFrame['DDe'].astype(float).round(4)
        compensateFrame["distance from max/min D single mutation"] = compensateFrame["distance from max/min D single mutation"].astype(float).round(4)

        flipFrame['Dm1'] = flipFrame['Dm1'].astype(float).round(4)
        flipFrame['Dm2'] = flipFrame['Dm2'].astype(float).round(4)
        flipFrame['Dm1m2'] = flipFrame['Dm1m2'].astype(float).round(4)
        flipFrame['DDe'] = flipFrame['DDe'].astype(float).round(4)
        flipFrame["distance from max/min D single mutation"] = flipFrame["distance from max/min D single mutation"].astype(float).round(4)


        frames = [rescueFrame,compensateFrame,antagFrame,flipFrame]

        for df in frames:
            df.loc[-1] = fields  # adding a row
            df.index = df.index + 1  # shifting index
            df = df.sort_index(inplace=True)  # sorting by index

        newDoc = Document()
        newDoc.add_heading(name,0)


        ##description paragraph
        paragraph1 = newDoc.add_paragraph(
            """rescue:     
            -list of sequences that have a rescue. Sorted by distance from DeltaE m1m2 to the highest DeltaE single mutation, which is the final column.
            -Columns: "Dm1","Dm2","Dm1m2","DDe","list of mutations","HD","DeltaE m1m2 - max(DeltaE m1, DeltaE m2)"""
        )
        for run in paragraph1.runs:
            font = run.font
            font.size= Pt(8)
        #adds table
        table = newDoc.add_table(rows=rescueFrame.shape[0],cols=rescueFrame.shape[1])
        table.allow_autofit = False
        table.autofit = False
        for i in range(rescueFrame.shape[0]):
            for j, cell in enumerate(table.rows[i].cells):
                cell.text = str(rescueFrame.values[i, j])
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(8)


        ##description paragraph
        paragraph2 = newDoc.add_paragraph(
            """compensate:
            -list of sequences that have a compensation. Sorted by distance from DeltaE m1m2 to min(DeltaE m1, DetlaE m2) which is the final col
            -Columns: "Dm1","Dm2","Dm1m2","DDe","list of mutations","HD","DeltaE m1m2 - min(DeltaE m1, DeltaE m2)"""
            
        )
        for run in paragraph2.runs:
            font = run.font
            font.size= Pt(8)
        #adds table
        table2 = newDoc.add_table(rows=compensateFrame.shape[0],cols=compensateFrame.shape[1])
        table2.allow_autofit = False
        table2.autofit = False
        for i in range(compensateFrame.shape[0]):
            for j, cell in enumerate(table2.rows[i].cells):
                cell.text = str(compensateFrame.values[i, j])
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(8)

                ##description paragraph
        paragraph3 = newDoc.add_paragraph(
            """
            antag.csv:
            -list of sequences that have a antagonistic interaction. Sorted by distance from DeltaE m1m2 to min(DeltaE m1, DetlaE m2) which is the final col
            -Columns: "Dm1","Dm2","Dm1m2","DDe","list of mutations","HD","DeltaE m1m2 - min(DeltaE m1, DeltaE m2)
            """
            
        )
        for run in paragraph3.runs:
            font = run.font
            font.size= Pt(8)
        #adds table
        table3 = newDoc.add_table(rows=antagFrame.shape[0],cols=antagFrame.shape[1])
        table3.allow_autofit = False
        table3.autofit = False
        for i in range(antagFrame.shape[0]):
            for j, cell in enumerate(table3.rows[i].cells):
                cell.text = str(antagFrame.values[i, j])
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(8)

                        ##description paragraph
        paragraph4 = newDoc.add_paragraph(
            """
            flipped.csv:
            -list of sequences that have m1/m2 flipped compared to the consensus. 
            -The interaction of these sequences are always the same as the consensus
                ie. if the consensus is a rescue, then the flipped.csv contains only rescues.
            -Columns: "Dm1","Dm2","Dm1m2","DDe","list of mutations","HD","DeltaE m1m2 - max/min(DeltaE m1, DeltaE m2)"
            """
            
        )
        for run in paragraph4.runs:
            font = run.font
            font.size= Pt(8)
        #adds table
        table4 = newDoc.add_table(rows=flipFrame.shape[0],cols=flipFrame.shape[1])
        table4.allow_autofit = False
        table4.autofit = False
        for i in range(flipFrame.shape[0]):
            for j, cell in enumerate(table4.rows[i].cells):
                cell.text = str(flipFrame.values[i, j])
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(8)
        newDoc.save(name+".docx")


    def writeToCSV(self,data,name):
        fields = ["Dm1","Dm2","Dm1m2","DDe","list of mutations","HD","distance from max/min D single mutation"]
        with open(name+".csv", "w") as outfile1:                
            write = csv.writer(outfile1)            
            write.writerow(fields)
            write.writerows(data)

    def writeToJson(self,data,name):
        with open(name+".json", "w") as outfile:
                json.dump(data, outfile)

        
        
        
        

            


    
    


        










        

        